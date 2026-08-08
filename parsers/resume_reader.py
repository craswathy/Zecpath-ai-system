import os
import pdfplumber
from docx import Document
from utils.logger import logger


def read_pdf(file_path):
    """Extract raw text from a PDF resume, page by page."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        logger.info(f"Successfully extracted PDF: {file_path}")
    except Exception as e:
        logger.error(f"Failed to extract PDF {file_path}: {e}")
    return text


def read_docx(file_path):
    """Extract raw text from a DOCX resume, including table cells."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        logger.info(f"Successfully extracted DOCX: {file_path}")
    except Exception as e:
        logger.error(f"Failed to extract DOCX {file_path}: {e}")
    return text


def extract_resume(file_path):
    """Route to correct reader based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_path}")
        return ""